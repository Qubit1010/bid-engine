"""Document ingestion: PDF (PyMuPDF) and DOCX (python-docx) -> pages of clean text."""
from pathlib import Path


def parse_pdf(path: Path) -> list[dict]:
    import fitz

    pages = []
    with fitz.open(path) as doc:
        for i, page in enumerate(doc):
            text = page.get_text("text")
            pages.append({"page": i + 1, "text": text.strip()})
    return pages


def parse_docx(path: Path) -> list[dict]:
    """DOCX has no fixed pages; emit pseudo-pages of ~3000 chars so source
    references stay meaningful."""
    import docx

    doc = docx.Document(path)
    blocks: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            blocks.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    pages, current, size = [], [], 0
    for b in blocks:
        current.append(b)
        size += len(b)
        if size >= 3000:
            pages.append({"page": len(pages) + 1, "text": "\n".join(current)})
            current, size = [], 0
    if current:
        pages.append({"page": len(pages) + 1, "text": "\n".join(current)})
    return pages


def parse_document(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        pages = parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        pages = parse_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix} (use PDF or DOCX)")

    non_empty = [p for p in pages if p["text"]]
    if not non_empty:
        raise ValueError("Document contains no extractable text (scanned image PDF?)")
    return {
        "num_pages": len(pages),
        "pages": pages,
        "chars": sum(len(p["text"]) for p in pages),
    }
