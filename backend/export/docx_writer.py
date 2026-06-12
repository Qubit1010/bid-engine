"""Export the reviewed proposal as a formatted DOCX + compliance matrix CSV."""
import csv
import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1D, 0x4E, 0xD8)
STATUS_COLORS = {"PASS": RGBColor(0x16, 0xA3, 0x4A), "PARTIAL": RGBColor(0xD9, 0x77, 0x06),
                 "GAP": RGBColor(0xDC, 0x26, 0x26)}


def build_proposal_docx(ws: dict, requirements: list[dict], sections: list[dict]) -> bytes:
    profile = ws.get("profile") or {}
    doc = Document()

    # title page
    title = doc.add_heading("Proposal Response", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(profile.get("title", ws["name"]))
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Issuer: {profile.get('issuer', '-')}\n"
        f"Sector: {profile.get('sector', '-')}\n"
        f"Submission deadline: {profile.get('submission_deadline') or '-'}\n"
        f"Prepared with BidSense on {date.today().isoformat()}"
    )
    winprob = ws.get("winprob") or {}
    if winprob:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            f"Bid decision: {winprob.get('decision', {}).get('decision', '-')} "
            f"(modeled win probability {winprob.get('probability', 0):.0%})"
        )
        r.bold = True
    doc.add_page_break()

    # sections (only approved + draft; reviewer controlled upstream)
    for s in sections:
        doc.add_heading(s["title"], level=1)
        status_p = doc.add_paragraph()
        sr = status_p.add_run(f"[{s['status'].upper()}]"
                              + (f"  Evidence: {', '.join(s.get('citations') or [])}"
                                 if s.get("citations") else ""))
        sr.font.size = Pt(8)
        sr.italic = True
        for para in (s.get("content") or "").split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # compliance matrix appendix
    doc.add_page_break()
    doc.add_heading("Appendix A — Compliance Matrix", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["#", "Requirement", "Mandatory", "Status", "Evidence"]):
        table.rows[0].cells[i].text = h
    for i, r in enumerate(requirements, 1):
        row = table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = r["text"][:300]
        row.cells[2].text = "Yes" if r["mandatory"] else "No"
        cell_p = row.cells[3].paragraphs[0]
        run = cell_p.add_run(r.get("status") or "-")
        if r.get("status") in STATUS_COLORS:
            run.font.color.rgb = STATUS_COLORS[r["status"]]
            run.bold = True
        used = r.get("used_cap_ids") or []
        if isinstance(used, str):
            used = []
        row.cells[4].text = ", ".join(used) if used else "-"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_compliance_csv(requirements: list[dict]) -> str:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["#", "Requirement", "Category", "Mandatory", "Status", "Confidence",
                "Rationale", "Evidence Cap IDs", "Source Page"])
    for i, r in enumerate(requirements, 1):
        used = r.get("used_cap_ids") or []
        w.writerow([
            i, r["text"], r.get("category", ""), "Yes" if r["mandatory"] else "No",
            r.get("status", ""), r.get("confidence", ""), r.get("rationale", ""),
            ", ".join(used) if isinstance(used, list) else used, r.get("source_page", ""),
        ])
    return buf.getvalue()
