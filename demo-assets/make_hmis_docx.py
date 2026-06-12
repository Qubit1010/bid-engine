"""Generate the third demo RFP as a DOCX (proves the DOCX ingestion path).

Engineered narrative: a strong Healthcare/IT bid -- high compliance, budget
near our won-bid median -- but ONE unclosable mandatory certification gap
(CMMI Level 5; the company profile holds Level 3). Expected outcome:
high P(win) + 1 mandatory gap => CONDITIONAL_GO ("win it if you close this").
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).parent / "RFP_Hospital_HMIS_Implementation.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)


def main():
    doc = Document()

    doc.add_heading("Request for Proposals (RFP)", 0)
    add_para(doc, "Implementation of an Integrated Hospital Management Information System (HMIS) "
                  "across Four Tertiary-Care Hospitals", bold=True)
    add_para(doc, "RFP Reference: PHD/HMIS/2026/07")
    add_para(doc, "Issued by: Provincial Health Department, Government of Khyber Pakhtunkhwa, Peshawar")
    add_para(doc, "Issue Date: 1 June 2026")

    add_heading(doc, "1. Introduction and Scope of Work")
    add_para(doc,
        "The Provincial Health Department (the Client) invites sealed proposals from qualified firms "
        "for the supply, implementation, integration, and support of an integrated Hospital Management "
        "Information System (HMIS) across four tertiary-care hospitals. The estimated total budget for "
        "this engagement is PKR 320 Million. The scope covers: (a) electronic medical records (EMR), "
        "outpatient and inpatient management; (b) laboratory and radiology information systems (LIS/RIS); "
        "(c) pharmacy and inventory management; (d) billing and insurance claims; (e) a central data "
        "warehouse with management dashboards; and (f) integration with existing hospital systems via "
        "HL7/FHIR standards. The successful bidder shall migrate legacy patient records, train hospital "
        "staff, and provide three (3) years of post-implementation support and maintenance.")

    add_heading(doc, "2. Mandatory Eligibility Requirements")
    add_para(doc, "Bidders failing any mandatory requirement below shall be disqualified.")
    for m in [
        "M1. The bidder must be registered with the Federal Board of Revenue (FBR) and appear on the "
        "Active Taxpayer List.",
        "M2. The bidder must hold a valid ISO 27001 information security management certification.",
        "M3. The bidder must hold a valid CMMI Level 5 appraisal for software development services.",
        "M4. The bidder must have successfully completed at least two (2) health-sector IT system "
        "implementations for hospitals or healthcare networks within the last five (5) years.",
        "M5. The assigned project manager must hold a valid PMP certification.",
        "M6. The bidder must submit a bid security of two percent (2%) of the bid price in the form of "
        "a bank guarantee or CDR from a scheduled bank.",
    ]:
        add_para(doc, m)

    add_heading(doc, "3. Technical Requirements")
    for t in [
        "T1. The HMIS shall support HL7 v2.x and FHIR R4 interoperability for exchange with existing "
        "laboratory analyzers and PACS systems.",
        "T2. The bidder must perform structured data migration of legacy patient records with a "
        "documented validation and reconciliation methodology.",
        "T3. The system shall provide role-based access control, audit trails, and encryption of "
        "patient data at rest and in transit.",
        "T4. The bidder must guarantee system availability of 99.5% measured monthly, with a disaster "
        "recovery site and a recovery time objective (RTO) of four (4) hours.",
        "T5. The bidder must deliver structured end-user and administrator training for approximately "
        "1,200 hospital staff across the four sites.",
        "T6. The bidder should provide mobile applications for physicians covering ward rounds, "
        "e-prescriptions, and result review.",
        "T7. The bidder must provide a help desk with 24/7 coverage and defined SLAs for incident "
        "response during the support period.",
    ]:
        add_para(doc, t)

    add_heading(doc, "4. Evaluation Criteria")
    add_para(doc, "Technical proposals shall be evaluated out of 100 marks as follows. Bidders scoring "
                  "less than 70 marks in the technical evaluation shall not be considered for financial "
                  "evaluation.")
    for c in [
        "Relevant health-sector experience and past performance — 30%",
        "Proposed technical solution and interoperability approach — 25%",
        "Implementation methodology, migration and training plan — 20%",
        "Key personnel qualifications — 15%",
        "Support model and SLAs — 10%",
    ]:
        add_para(doc, c)

    add_heading(doc, "5. Submission Requirements and Key Dates")
    for s in [
        "Proposals shall be submitted in two separate sealed envelopes marked \"Technical Proposal\" "
        "and \"Financial Proposal\".",
        "Pre-bid meeting: 18 June 2026 at 11:00 AM at the Provincial Health Department, Peshawar.",
        "Deadline for submission of written queries: 22 June 2026.",
        "Proposal submission deadline: 6 July 2026 at 2:00 PM local time.",
        "Technical bid opening: 6 July 2026 at 2:30 PM.",
        "The bid validity period shall be 90 days from the submission deadline.",
        "The bidder must respond to the required proposal sections Q1 to Q5 listed below.",
    ]:
        add_para(doc, s)

    add_heading(doc, "6. Required Proposal Sections (Q1-Q5)")
    for q in [
        "Q1. Company profile, registrations, and certifications.",
        "Q2. Relevant project experience with verifiable references.",
        "Q3. Proposed solution architecture and interoperability approach.",
        "Q4. Implementation plan, migration methodology, and training plan.",
        "Q5. Support and maintenance model with SLAs.",
    ]:
        add_para(doc, q)

    doc.save(OUT)
    print(f"saved {OUT}")


if __name__ == "__main__":
    main()
